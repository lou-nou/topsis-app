import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
from io import BytesIO
from PIL import Image
from fpdf import FPDF
import tempfile
from datetime import datetime
import pdfkit
from tempfile import NamedTemporaryFile


# --- Affichage du logo ---
try:
    logo = Image.open("logo.png")
    st.image(logo, width=120)
except FileNotFoundError:
    pass

st.title("🔍 Analyse multicritère TOPSIS avec sensibilité")

# --- Upload du fichier Excel ---
uploaded_file = st.file_uploader(
    "📤 Charger un fichier Excel (.xlsx) contenant les feuilles 'données' et 'ponderation'",
    type=["xlsx"]
)

if uploaded_file is None:
    st.info("Veuillez importer un fichier pour commencer.")
    st.stop()

# --- Lecture des données ---
try:
    df_data = pd.read_excel(uploaded_file, sheet_name='données', engine='openpyxl')
    df_weights = pd.read_excel(uploaded_file, sheet_name='ponderation', engine='openpyxl')
except Exception as e:
    st.error(f"Erreur de lecture du fichier : {e}")
    st.stop()

# --- Préparation des données ---
actions = df_data.iloc[:, 0]
criteria_data = df_data.iloc[:, 1:].apply(pd.to_numeric, errors='coerce')
criteria_names = criteria_data.columns

# --- Normalisation des poids ---
weights_series = df_weights.set_index(df_weights.columns[0]).iloc[:, 0]
weights_series = weights_series.reindex(criteria_names).fillna(0)
weights_series = weights_series / weights_series.sum()

# --- Détermination des impacts ---
impacts = ['min' if ('coût' in crit.lower() or 'roi' in crit.lower() or
                     'facilité' in crit.lower() or 'acceptabilité' in crit.lower()) else 'max'
           for crit in criteria_names]

# --- Fonction TOPSIS ---
def topsis(criteria_data, weights, impacts):
    norm_matrix = criteria_data / np.sqrt((criteria_data ** 2).sum())
    weighted_matrix = norm_matrix * weights

    ideal_best = weighted_matrix.max()
    ideal_worst = weighted_matrix.min()

    for i, impact in enumerate(impacts):
        if impact == 'min':
            ideal_best.iloc[i], ideal_worst.iloc[i] = ideal_worst.iloc[i], ideal_best.iloc[i]

    dist_best = np.sqrt(((weighted_matrix - ideal_best) ** 2).sum(axis=1))
    dist_worst = np.sqrt(((weighted_matrix - ideal_worst) ** 2).sum(axis=1))
    scores = dist_worst / (dist_best + dist_worst)
    return scores

# --- Calcul TOPSIS ---
scores = topsis(criteria_data, weights_series, impacts)
df_result = pd.DataFrame({
    'Action': actions,
    'Score TOPSIS': scores
})
df_result['Classement'] = df_result['Score TOPSIS'].rank(ascending=False).astype(int)
df_result = df_result.sort_values(by='Score TOPSIS', ascending=False)

# --- Analyse de sensibilité ---
excluded_criterion = "Économie d’énergie (%)"
sensitivity_results = []

for i, crit in enumerate(criteria_names):
    if crit != excluded_criterion:
        for variation in [-0.1, 0.1]:
            new_weights = weights_series.copy()
            new_weights.iloc[i] *= (1 + variation)
            new_weights /= new_weights.sum()
            new_scores = topsis(criteria_data, new_weights, impacts)
            for action, score in zip(actions, new_scores):
                sensitivity_results.append({
                    'Critère modifié': crit,
                    'Variation': f"{int(variation * 100)}%",
                    'Action': action,
                    'Score TOPSIS': score
                })

df_sensitivity = pd.DataFrame(sensitivity_results)

# --- Graphiques ---
fig_bar = px.bar(df_result, x='Action', y='Score TOPSIS', color='Classement', text='Classement',
                 title="Scores TOPSIS par action")

selected_criteria = st.multiselect(
    "Filtrer les critères affichés dans l’analyse de sensibilité",
    options=df_sensitivity['Critère modifié'].unique(),
    default=list(df_sensitivity['Critère modifié'].unique())
)
filtered_df = df_sensitivity[df_sensitivity['Critère modifié'].isin(selected_criteria)]

fig_sens = px.line(filtered_df, x='Action', y='Score TOPSIS',
                   color='Critère modifié', line_dash='Variation', markers=True,
                   title="Variation des scores TOPSIS selon les poids des critères")
fig_sens.update_layout(yaxis=dict(range=[0, 1]), height=700)

# --- Affichage dans Streamlit ---
st.subheader("📊 Résultats de l’analyse TOPSIS")
st.dataframe(df_result, use_container_width=True)
st.plotly_chart(fig_bar, use_container_width=True)

st.subheader("📉 Analyse de sensibilité")
st.plotly_chart(fig_sens, use_container_width=True)


#---Creation du pdf
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from io import BytesIO
import plotly.io as pio
from datetime import datetime
import base64

def generate_word_direct(df_result, selected_criteria, fig_bar, fig_sens):
    """Génère un Word sans utiliser de fichiers temporaires"""
    
    # 1. Initialisation du document
    doc = Document()
    
    # 2. Style par défaut
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # 3. Contenu du rapport
    doc.add_heading('Rapport TOPSIS', 0)
    doc.add_paragraph(f"Généré le {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    
    # 4. Tableau des résultats
    doc.add_heading('Résultats', level=1)
    table = doc.add_table(rows=1, cols=len(df_result.columns), style="Light Shading")
    
    # En-têtes
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df_result.columns):
        hdr_cells[i].text = str(col)
    
    # Données
    for _, row in df_result.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # 5. Graphiques (méthode directe)
    def add_plot(fig, title):
        """Ajoute un graphique directement depuis la mémoire"""
        img_bytes = fig.to_image(format="png", width=800, height=500)
        doc.add_heading(title, level=2)
        doc.add_picture(BytesIO(img_bytes), width=Inches(6))
    
    # Ajout des graphiques
    add_plot(fig_bar, "Scores TOPSIS")
    add_plot(fig_sens, "Analyse de sensibilité")
    
    # 6. Sauvegarde en mémoire
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Bouton d'export
if st.button("📄 Exporter Word (mémoire)"):
    try:
        # Vérification de kaleido
        if not pio.kaleido.scope:
            st.warning("Installation de Kaleido...")
            import subprocess
            import sys
            subprocess.check_call([sys.executable, "-m", "pip", "install", "kaleido"])
            import plotly.io as pio
            
        docx_file = generate_word_direct(df_result, selected_criteria, fig_bar, fig_sens)
        
        st.download_button(
            label="💾 Télécharger DOCX",
            data=docx_file,
            file_name="rapport_topsis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Erreur: {str(e)}")